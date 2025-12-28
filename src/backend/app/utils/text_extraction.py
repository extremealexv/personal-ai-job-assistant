"""Text extraction utilities for resume parsing."""
from pathlib import Path
from typing import Optional

import PyPDF2
from docx import Document
from fastapi import HTTPException, status


class TextExtractor:
    """Extract text from PDF and DOCX files."""

    @staticmethod
    async def extract_from_pdf(file_path: Path) -> str:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file.

        Returns:
            Extracted text content.

        Raises:
            HTTPException: If PDF cannot be read.
        """
        try:
            text_parts = []

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as page_error:
                        # Log error but continue with other pages
                        print(f"Error extracting page {page_num + 1}: {page_error}")
                        continue

            if not text_parts:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Could not extract text from PDF. File may be empty or corrupted.",
                )

            # Join all pages with double newline
            return "\n\n".join(text_parts)

        except PyPDF2.errors.PdfReadError as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid PDF file: {str(e)}",
            ) from e
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error processing PDF: {str(e)}",
            ) from e

    @staticmethod
    async def extract_from_docx(file_path: Path) -> str:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file.

        Returns:
            Extracted text content.

        Raises:
            HTTPException: If DOCX cannot be read.
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text)
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))

            if not text_parts:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Could not extract text from DOCX. File may be empty or corrupted.",
                )

            # Join all parts with single newline
            return "\n".join(text_parts)

        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error processing DOCX: {str(e)}",
            ) from e

    @classmethod
    async def extract_text(cls, file_path: Path) -> str:
        """Extract text from file based on extension.

        Args:
            file_path: Path to file.

        Returns:
            Extracted text content.

        Raises:
            HTTPException: If file type is not supported or extraction fails.
        """
        file_ext = file_path.suffix.lower()

        if file_ext == ".pdf":
            return await cls.extract_from_pdf(file_path)
        elif file_ext == ".docx":
            return await cls.extract_from_docx(file_path)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {file_ext}",
            )
